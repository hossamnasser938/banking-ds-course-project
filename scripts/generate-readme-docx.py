#!/usr/bin/env python3
"""Convert README.md to a Word document."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_paragraph(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    return paragraph


def add_code_block(doc: Document, lines: list[str]) -> None:
    text = "\n".join(lines)
    if text.strip().startswith("flowchart") or text.strip().startswith("graph"):
        doc.add_paragraph(
            "Architecture diagram: User/Browser → Edge Gateway → Authentication Service, "
            "Observability Service, and Banking Load Balancer → Banking Replicas (Zone A and Zone B) → PostgreSQL."
        )
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)


def parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if not line.strip().startswith("|"):
            break
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
            if r_idx == 0:
                set_cell_shading(cell, "E8E8E8")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def convert_markdown_to_docx(markdown_path: Path, output_path: Path) -> None:
    content = markdown_path.read_text(encoding="utf-8")
    lines = content.splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table_rows(table_lines))
            continue

        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            add_formatted_paragraph(doc, text, style="List Number")
            i += 1
            continue

        if stripped.startswith("- "):
            add_formatted_paragraph(doc, stripped[2:], style="List Bullet")
            i += 1
            continue

        if stripped:
            add_formatted_paragraph(doc, stripped)
        i += 1

    doc.save(output_path)


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    markdown_path = root / "README.md"
    output_path = root / "Distributed-Banking-System-Documentation.docx"

    if not markdown_path.exists():
        print(f"Missing README: {markdown_path}", file=sys.stderr)
        return 1

    convert_markdown_to_docx(markdown_path, output_path)
    print(f"Created: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
